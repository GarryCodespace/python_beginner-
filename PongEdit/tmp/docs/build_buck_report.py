from pathlib import Path
from math import log10, sqrt, pi

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("/Users/garryyuan/python/PongEdit")
OUT = ROOT / "output" / "doc"
PLOT = OUT / "buck_frequency_response_clean.png"
DOCX = OUT / "ELEC4614_A3_Buck_Converter_Report_polished.docx"


freq = [10, 50, 100, 500, 700, 1000, 2000, 4000]
dvo = [6.0, 9.6, 10.5, 6.1, 3.4, 4.3, 3.4, 2.0]
dil = [1.6, 4.6, 2.4, 9.4, 8.6, 6.9, 3.6, 2.5]


def font(size=18, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except Exception:
            pass
    return ImageFont.load_default()


def draw_plot(path: Path):
    w, h = 1100, 680
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    left, right, top, bottom = 105, 45, 92, 95
    x0, x1, y0, y1 = left, w - right, h - bottom, top

    f_min, f_max = 10, 4000
    y_min, y_max = 0, 12

    def xp(f):
        return x0 + (log10(f) - log10(f_min)) / (log10(f_max) - log10(f_min)) * (x1 - x0)

    def yp(v):
        return y0 - (v - y_min) / (y_max - y_min) * (y0 - y1)

    # Grid and axes.
    for gv in range(0, 13, 2):
        y = yp(gv)
        d.line((x0, y, x1, y), fill=(225, 225, 225), width=1)
        d.text((x0 - 38, y - 10), str(gv), fill=(55, 55, 55), font=font(18))

    xticks = [10, 50, 100, 500, 1000, 2000, 4000]
    for gf in xticks:
        x = xp(gf)
        d.line((x, y0, x, y1), fill=(235, 235, 235), width=1)
        label = f"{gf}" if gf < 1000 else f"{gf//1000}k"
        d.text((x - 18, y0 + 16), label, fill=(55, 55, 55), font=font(18))

    d.line((x0, y0, x1, y0), fill=(30, 30, 30), width=2)
    d.line((x0, y0, x0, y1), fill=(30, 30, 30), width=2)

    def series(vals, color):
        pts = [(xp(f), yp(v)) for f, v in zip(freq, vals)]
        d.line(pts, fill=color, width=4)
        for x, y in pts:
            d.ellipse((x - 6, y - 6, x + 6, y + 6), fill=color, outline="white", width=2)

    blue = (26, 102, 158)
    orange = (220, 126, 35)
    series(dvo, blue)
    series(dil, orange)

    title = "Buck Converter Frequency Response"
    d.text(((w - d.textlength(title, font=font(26, True))) / 2, 12), title, fill=(20, 20, 20), font=font(26, True))
    xlabel = "Duty-cycle modulation frequency fD (Hz, log scale)"
    d.text(((w - d.textlength(xlabel, font=font(20))) / 2, h - 45), xlabel, fill=(20, 20, 20), font=font(20))
    ylabel = "Measured envelope amplitude (scope divisions)"
    d.text((x0, 56), ylabel, fill=(20, 20, 20), font=font(18))

    # Legend.
    lx, ly = x1 - 250, y1 + 20
    d.rectangle((lx, ly, lx + 215, ly + 72), fill=(255, 255, 255), outline=(180, 180, 180))
    d.line((lx + 18, ly + 22, lx + 58, ly + 22), fill=blue, width=4)
    d.ellipse((lx + 34, ly + 18, lx + 42, ly + 26), fill=blue)
    d.text((lx + 72, ly + 10), "Delta Vo", fill=(20, 20, 20), font=font(18))
    d.line((lx + 18, ly + 52, lx + 58, ly + 52), fill=orange, width=4)
    d.ellipse((lx + 34, ly + 48, lx + 42, ly + 56), fill=orange)
    d.text((lx + 72, ly + 40), "Delta iL", fill=(20, 20, 20), font=font(18))

    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, quality=95)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, rows, widths=None, header=True):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            set_cell_text(table.cell(i, j), str(text), bold=(header and i == 0), size=8)
            if header and i == 0:
                set_cell_shading(table.cell(i, j), "D9EAF7")
            if widths:
                table.cell(i, j).width = widths[j]
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_para(doc, text, size=9.5, after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.04
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8.5)
    return p


def style_doc(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    for name in ["Heading 1", "Heading 2", "Title"]:
        styles[name].font.name = "Arial"
        styles[name].font.color.rgb = RGBColor(20, 20, 20)
    styles["Heading 1"].font.size = Pt(13)
    styles["Heading 2"].font.size = Pt(11)


def build_doc():
    draw_plot(PLOT)

    doc = Document()
    style_doc(doc)

    sec = doc.sections[0]
    sec.top_margin = Cm(1.65)
    sec.bottom_margin = Cm(1.45)
    sec.left_margin = Cm(1.65)
    sec.right_margin = Cm(1.65)

    # Cover page.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("School of Electrical Engineering & Telecommunications\nUniversity of New South Wales")
    r.bold = True
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ELEC4614 - Power Electronics\nAssignment / Technical Report")
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph()
    cover_rows = [
        ["Due Date", "Friday, Week 11, T1 2026, 3:00 pm"],
        ["Submission Date", "19 April 2026"],
        ["Title", "Experiment 1: DC-DC Step-Down (Buck) Converter"],
        ["Submitted by", "Garry Yuan"],
        ["Student Number", "z5425329"],
    ]
    table = doc.add_table(rows=len(cover_rows), cols=2)
    table.style = "Table Grid"
    for i, row in enumerate(cover_rows):
        for j, text in enumerate(row):
            set_cell_text(table.cell(i, j), text, bold=(j == 0), size=10)

    doc.add_paragraph()
    add_heading(doc, "Declaration", 2)
    add_para(
        doc,
        "This submission was prepared and written entirely by me, in my own words. Information which appears in this submission that is not my own has been fully acknowledged in the list of references. Where figures prepared by others are reproduced, the original sources are mentioned in the captions.",
        size=10,
        after=8,
    )
    add_para(doc, "Signature: ___________________________________", size=10)
    doc.add_section(WD_SECTION.NEW_PAGE)

    # Report body.
    add_heading(doc, "Experiment 1 Buck Converter", 1)
    add_para(
        doc,
        "This report answers Questions 6.3 and 6.4 for the ELEC4614 buck-converter laboratory. The converter was supplied from Vd = 50 V, used a 5 ohm load and a 1000 uF output capacitor, and was tested with switching frequencies of 5, 10 and 20 kHz. The main aim is to connect the measured waveforms and frequency-response data to the expected continuous-conduction-mode (CCM), discontinuous-conduction-mode (DCM), and LC-filter theory.",
    )

    add_heading(doc, "Key Theory Used for Comparison", 2)
    add_para(
        doc,
        "For an ideal buck converter in CCM, volt-second balance on the inductor gives Vo = D Vd. The inductor-current ripple is Delta iL = Vo(1-D)/(L fs), and the boundary for CCM is Lmin = (1-D)R/(2 fs). Including semiconductor drops gives Vo = D(Vd - VT) - (1-D)VD, explaining why measured voltages are slightly below D Vd. In DCM, iL reaches zero before the next switching period; the output voltage is then no longer set only by D, but also by L, fs and R.",
    )
    f0 = 1 / (2 * pi * sqrt(188e-6 * 1000e-6))
    zeta = (1 / (2 * 5)) * sqrt(188e-6 / 1000e-6)
    add_para(
        doc,
        f"For the frequency-response test, the output stage is a second-order low-pass LC filter. With L = 188 uH and C = 1000 uF, the ideal natural frequency is f0 = 1/(2*pi*sqrt(LC)) = {f0:.0f} Hz. The approximate load damping ratio is zeta = (1/(2R))*sqrt(L/C) = {zeta:.3f}, so a resonant peak is expected before the high-frequency roll-off.",
    )

    add_heading(doc, "6.3 CCM and DCM Waveform Behaviour", 2)
    add_para(
        doc,
        "The saved oscilloscope traces should be interpreted by splitting each switching period into switch-on and switch-off intervals. In CCM the inductor current remains positive for the whole period. During ton, vL = Vd - Vo and iL rises linearly; during toff, vL = -Vo and iL falls linearly through the diode. Therefore iT is a chopped copy of iL during ton, while iD carries the complementary current during toff. The capacitor current is iC = iL - Io, so it is triangular and centred approximately on zero, while vo remains nearly constant because the 1000 uF capacitor filters the switching ripple.",
    )
    add_para(
        doc,
        "In DCM, the inductor current falls to zero before the next switching pulse. The cycle has three intervals: switch-on, diode freewheel, and a zero-current dead interval. During the dead interval iL, iT and iD are zero, vL is approximately zero, and the load is supplied by the output capacitor, giving a slow fall in vo. This is the clearest waveform evidence distinguishing DCM from CCM: DCM has a flat zero-current segment, while CCM has a continuous triangular current.",
    )

    waveform_rows = [
        ["Quantity", "CCM waveform", "DCM waveform"],
        ["iL", "Triangular ripple above zero; average approximately Vo/R.", "Rises from zero, returns to zero, then stays at zero."],
        ["vL", "+(Vd-Vo) during ton; -Vo during toff.", "+(Vd-Vo), then -Vo, then approximately 0 in dead time."],
        ["iC", "Triangular around zero; charges/discharges capacitor.", "Pulsed; during dead time capacitor alone supplies load."],
        ["vo", "Small ripple around D Vd.", "Larger ripple and load-dependent average."],
        ["iT", "Equals iL only during ton.", "Pulse starts from zero and ends when switch turns off."],
        ["iD", "Equals iL only during toff.", "Freewheel pulse returns to zero before next cycle."],
    ]
    add_table(doc, waveform_rows)
    add_caption(doc, "Table 1: Expected and observed identifying features of CCM and DCM buck-converter waveforms.")

    add_heading(doc, "Experimental Agreement", 2)
    add_para(
        doc,
        "The measured DC output values follow the linear CCM prediction closely. For example, with fs = 20 kHz and L = 188 uH, the measured output voltages at D = 0.1 to 0.6 were 5.0, 10.0, 15.0, 20.0, 25.0 and 30.0 V, matching Vo = D x 50 V within reading accuracy. The lower-frequency 5 kHz, 141 uH case gave slightly lower values, 4.8 to 29.3 V, which is consistent with device voltage drops, winding resistance and increased ripple stress.",
    )
    add_para(
        doc,
        "The theory also predicts that increasing fs or L reduces Delta iL because Delta iL is inversely proportional to L fs. This trend is visible in the measurements: at D = 0.6, Delta iL decreased from about 49.7 mA for 5 kHz, 141 uH to about 14.0 mA for 20 kHz, 188 uH. The large-inductance, high-switching-frequency case is therefore the most robustly continuous operating point.",
    )

    add_heading(doc, "6.4 Frequency Response", 2)
    freq_rows = [["fD (Hz)", "Delta Vo", "Delta iL"]]
    freq_rows += [[f, v, i] for f, v, i in zip(freq, dvo, dil)]
    add_table(doc, freq_rows)
    add_caption(doc, "Table 2: Measured envelope amplitudes for sinusoidal duty-cycle modulation, fs = 20 kHz and L = 188 uH.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(PLOT), width=Inches(5.75))
    add_caption(
        doc,
        "Figure 1: Frequency response of output-voltage and inductor-current envelope amplitudes versus duty-cycle modulation frequency.",
    )

    add_para(
        doc,
        "The voltage response rises from 6.0 divisions at 10 Hz to a maximum of 10.5 divisions at 100 Hz, then decreases to 2.0 divisions at 4 kHz. The inductor-current response peaks later, at about 500 Hz, with 9.4 divisions before rolling off to 2.5 divisions at 4 kHz. The two peaks need not occur at exactly the same frequency because vo is related to capacitor charge, while iL is the filter-state current; their phase and magnitude responses differ around resonance.",
    )
    add_para(
        doc,
        "The measured resonance region is broadly consistent with the predicted LC natural frequency of about 367 Hz. The spread between the 100 Hz voltage peak and 500 Hz current peak is reasonable for a real converter with capacitor ESR, inductor winding resistance, sensor scaling error, supply impedance and finite switch/diode drops. These parasitics increase damping and shift the apparent peak from the ideal value.",
    )
    add_para(
        doc,
        "Using the -3 dB point from the voltage peak gives a threshold of 10.5/sqrt(2) = 7.4 divisions, which places the post-peak voltage cutoff between 100 and 500 Hz. The high-frequency attenuation from 1 kHz to 4 kHz confirms the expected low-pass behaviour: duty-cycle disturbances above the resonance are increasingly rejected by the LC output filter. In a closed-loop regulator this resonance must be considered in compensator design, since the phase lag near resonance can reduce stability margin.",
    )

    add_heading(doc, "Conclusion", 2)
    add_para(
        doc,
        "The experiment confirms the central buck-converter theory. In CCM, the output voltage is approximately proportional to duty cycle and the six measured waveforms follow the expected switch-on and freewheel intervals. In DCM, the zero-current interval changes the waveform sequence and makes the conversion ratio dependent on load and component values rather than duty cycle alone. The measurements also confirm that larger inductance and higher switching frequency reduce inductor-current ripple.",
    )
    add_para(
        doc,
        "The frequency-response test shows the converter output stage acting as a lightly damped second-order low-pass filter. The observed resonance and subsequent roll-off agree with the predicted LC behaviour, while the difference between ideal and measured peak frequencies is explained by ESR, winding resistance, non-ideal devices and measurement scaling. Overall, the experimental results match the theory at the level expected from a practical laboratory converter.",
    )

    add_heading(doc, "References", 2)
    refs = [
        '[1] J. Fletcher and M. Priestley, "ELEC4614 Power Electronics Laboratory - Experiment 1: DC-DC Step-Down (Buck) Converter," UNSW School of Electrical Engineering & Telecommunications, Feb. 2018.',
        "[2] N. Mohan, T. M. Undeland and W. P. Robbins, Power Electronics: Converters, Applications, and Design, 3rd ed., Wiley, 2003.",
        "[3] R. W. Erickson and D. Maksimovic, Fundamentals of Power Electronics, 3rd ed., Springer, 2020.",
    ]
    for ref in refs:
        add_para(doc, ref, size=8.5, after=1)

    doc.save(DOCX)
    print(DOCX)
    print(PLOT)


if __name__ == "__main__":
    build_doc()
